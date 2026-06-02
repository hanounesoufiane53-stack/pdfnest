from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os, uuid, zipfile, hashlib, secrets
from pathlib import Path
from werkzeug.utils import secure_filename
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
import urllib.request
import json

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = Path("uploads")
OUTPUT_FOLDER = Path("outputs")
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

# ── SUPABASE CONFIG ──
SUPABASE_URL = "https://vglbfdmamherpnihmjxu.supabase.co"
SUPABASE_KEY = "sb_publishable_ZXfEsoQZ34ZcR76yxXUq0g_Dk4HzWUc"
SUPABASE_SECRET = "sb_secret_oDrfLkOpabRKjBaEEz063g_aFyldcpM"

def supabase_request(method, endpoint, data=None):
    url = f"{SUPABASE_URL}/rest/v1/{endpoint}"
    headers = {
        "apikey": SUPABASE_SECRET,
        "Authorization": f"Bearer {SUPABASE_SECRET}",
        "Content-Type": "application/json",
        "Prefer": "return=representation"
    }
    body = json.dumps(data).encode() if data else None
    req = urllib.request.Request(url, data=body, headers=headers, method=method)
    try:
        with urllib.request.urlopen(req) as res:
            return json.loads(res.read().decode())
    except urllib.error.HTTPError as e:
        error_body = e.read().decode()
        print(f"Supabase error: {error_body}")
        return {"error": error_body}
    except Exception as e:
        print(f"Request error: {e}")
        return {"error": str(e)}

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def generate_token(user_id, email):
    return hashlib.sha256(f"{user_id}{email}{secrets.token_hex(16)}".encode()).hexdigest()

# Simple in-memory session store
sessions = {}

def track_tool_use(tool_name, user_email=None, file_size=0):
    try:
        # Record tool usage
        supabase_request('POST', 'analytics', {
            "tool_name": tool_name,
            "user_email": user_email,
            "file_size": file_size
        })
        # Update total files count
        stats = supabase_request('GET', 'site_stats?select=id,total_files')
        if stats and len(stats) > 0:
            new_count = stats[0]['total_files'] + 1
            supabase_request('PATCH', f'site_stats?id=eq.{stats[0]["id"]}', {
                "total_files": new_count,
                "updated_at": "now()"
            })
    except Exception as e:
        print(f"Analytics error: {e}")

def get_real_stats():
    try:
        stats = supabase_request('GET', 'site_stats?select=total_files,total_users')
        users_count = supabase_request('GET', 'users?select=id')
        tool_usage = supabase_request('GET', 'analytics?select=tool_name')
        total_files = stats[0]['total_files'] if stats and len(stats) > 0 else 0
        total_users = len(users_count) if users_count else 0
        # Count tool usage
        tool_counts = {}
        if tool_usage:
            for item in tool_usage:
                t = item['tool_name']
                tool_counts[t] = tool_counts.get(t, 0) + 1
        return {
            "total_files": total_files,
            "total_users": total_users,
            "tool_counts": tool_counts
        }
    except Exception as e:
        print(f"Stats error: {e}")
        return {"total_files": 0, "total_users": 0, "tool_counts": {}}


def save_upload(file):
    uid = str(uuid.uuid4())
    filename = secure_filename(file.filename)
    path = UPLOAD_FOLDER / f"{uid}_{filename}"
    file.save(path)
    return path

def out(name):
    return OUTPUT_FOLDER / f"{str(uuid.uuid4())}_{name}"


def track_tool_use(tool_name, user_email=None, file_size=0):
    try:
        # Log tool usage
        supabase_request('POST', 'analytics', {
            "tool_name": tool_name,
            "user_email": user_email,
            "file_size": file_size
        })
        # Update total files count
        stats = supabase_request('GET', 'site_stats?select=id,total_files')
        if stats and len(stats) > 0:
            new_count = stats[0]['total_files'] + 1
            supabase_request('PATCH', f'site_stats?id=eq.{stats[0]["id"]}', {
                "total_files": new_count,
                "updated_at": "now()"
            })
    except Exception as e:
        print(f"Analytics error: {e}")

def track_new_user():
    try:
        stats = supabase_request('GET', 'site_stats?select=id,total_users')
        if stats and len(stats) > 0:
            new_count = stats[0]['total_users'] + 1
            supabase_request('PATCH', f'site_stats?id=eq.{stats[0]["id"]}', {
                "total_users": new_count,
                "updated_at": "now()"
            })
    except Exception as e:
        print(f"Analytics error: {e}")

# ── SERVE FRONTEND ──
@app.route('/')
def home():
    return send_file('index.html')

@app.route('/api/status')
def status():
    return jsonify({"status": "PDFnest is running!"})

# ── AUTH ROUTES ──
@app.route('/api/auth/signup', methods=['POST'])
def signup():
    try:
        data = request.json
        name = data.get('name','').strip()
        email = data.get('email','').strip().lower()
        password = data.get('password','')
        if not name or not email or not password:
            return jsonify({"error": "All fields required"}), 400
        if len(password) < 6:
            return jsonify({"error": "Password must be at least 6 characters"}), 400
        # Check if email exists
        existing = supabase_request('GET', f'users?email=eq.{email}&select=id')
        if existing and len(existing) > 0:
            return jsonify({"error": "Email already registered"}), 400
        # Create user
        hashed = hash_password(password)
        new_user = supabase_request('POST', 'users', {
            "name": name,
            "email": email,
            "password": hashed,
            "is_admin": False
        })
        if isinstance(new_user, list) and len(new_user) > 0:
            user = new_user[0]
            token = generate_token(user['id'], email)
            sessions[token] = {"id": user['id'], "name": user['name'], "email": email, "is_admin": False}
            track_new_user()
            return jsonify({"token": token, "name": user['name'], "email": email, "is_admin": False})
        return jsonify({"error": "Signup failed"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/auth/login', methods=['POST'])
def login():
    try:
        data = request.json
        email = data.get('email','').strip().lower()
        password = data.get('password','')
        if not email or not password:
            return jsonify({"error": "Email and password required"}), 400
        hashed = hash_password(password)
        users = supabase_request('GET', f'users?email=eq.{email}&select=id,name,email,is_admin,password')
        if not users or len(users) == 0:
            return jsonify({"error": "Invalid email or password"}), 401
        user = users[0]
        # Check admin with plain password too for initial setup
        if user['password'] != hashed and user['password'] != password:
            return jsonify({"error": "Invalid email or password"}), 401
        token = generate_token(user['id'], email)
        sessions[token] = {"id": user['id'], "name": user['name'], "email": email, "is_admin": user['is_admin']}
        return jsonify({"token": token, "name": user['name'], "email": email, "is_admin": user['is_admin']})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/auth/me', methods=['GET'])
def me():
    token = request.headers.get('Authorization','').replace('Bearer ','')
    if token in sessions:
        return jsonify(sessions[token])
    return jsonify({"error": "Not authenticated"}), 401

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    token = request.headers.get('Authorization','').replace('Bearer ','')
    sessions.pop(token, None)
    return jsonify({"success": True})


@app.route('/api/analytics', methods=['GET'])
def get_analytics():
    try:
        stats = supabase_request('GET', 'site_stats?select=*')
        tools = supabase_request('GET', 'analytics?select=tool_name&order=created_at.desc&limit=100')
        users_count = supabase_request('GET', 'users?select=id')
        
        # Count tool usage
        tool_counts = {}
        if isinstance(tools, list):
            for t in tools:
                name = t.get('tool_name','')
                tool_counts[name] = tool_counts.get(name, 0) + 1
        
        total_files = 0
        total_users = 0
        if isinstance(stats, list) and len(stats) > 0:
            total_files = stats[0].get('total_files', 0)
        if isinstance(users_count, list):
            total_users = len(users_count)
            
        return jsonify({
            "total_files": total_files,
            "total_users": total_users,
            "tool_counts": tool_counts
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ── PDF TOOLS ──
@app.route('/api/merge', methods=['POST'])
def merge_pdf():
    try:
        files = request.files.getlist('files')
        if len(files) < 2:
            return jsonify({"error": "Upload at least 2 PDFs"}), 400
        merger = PdfMerger()
        saved = []
        for f in files:
            p = save_upload(f); saved.append(p)
            merger.append(str(p))
        o = out("merged.pdf")
        merger.write(str(o)); merger.close()
        for p in saved: p.unlink(missing_ok=True)
        track_tool_use('Merge PDF')
        track_tool_use('merge')
        return send_file(o, as_attachment=True, download_name="pdfnest-merged.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/split', methods=['POST'])
def split_pdf():
    try:
        file = request.files.get('file')
        page_range = request.form.get('range', '')
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        total = len(reader.pages)
        pages = set()
        if page_range:
            for part in page_range.split(','):
                part = part.strip()
                if '-' in part:
                    a,b = part.split('-')
                    pages.update(range(int(a)-1, int(b)))
                else:
                    pages.add(int(part)-1)
        else:
            pages = set(range(total))
        zip_out = out("split.zip")
        with zipfile.ZipFile(zip_out, 'w') as zf:
            for i in sorted(pages):
                if 0 <= i < total:
                    w = PdfWriter()
                    w.add_page(reader.pages[i])
                    pf = out(f"page_{i+1}.pdf")
                    with open(pf, 'wb') as f2: w.write(f2)
                    zf.write(pf, f"page_{i+1}.pdf")
                    pf.unlink(missing_ok=True)
        path.unlink(missing_ok=True)
        track_tool_use('Split PDF')
        track_tool_use('split')
        return send_file(zip_out, as_attachment=True, download_name="pdfnest-split.zip")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compress', methods=['POST'])
def compress_pdf():
    try:
        file = request.files.get('file')
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        writer = PdfWriter()
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        o = out("compressed.pdf")
        with open(o, 'wb') as f: writer.write(f)
        path.unlink(missing_ok=True)
        track_tool_use('Compress PDF')
        track_tool_use('compress')
        return send_file(o, as_attachment=True, download_name="pdfnest-compressed.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf2word', methods=['POST'])
def pdf_to_word():
    try:
        from docx import Document
        file = request.files.get('file')
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        doc = Document()
        doc.add_heading('Converted PDF', 0)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                doc.add_heading(f'Page {i+1}', level=1)
                doc.add_paragraph(text)
        o = out("converted.docx")
        doc.save(str(o))
        path.unlink(missing_ok=True)
        track_tool_use('PDF to Word')
        track_tool_use('pdf2word')
        return send_file(o, as_attachment=True, download_name="pdfnest-converted.docx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/pdf2jpg', methods=['POST'])
def pdf_to_jpg():
    try:
        from PIL import Image, ImageDraw
        file = request.files.get('file')
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        zip_out = out("images.zip")
        with zipfile.ZipFile(zip_out, 'w') as zf:
            for i, page in enumerate(reader.pages):
                img = Image.new('RGB', (800, 1100), color='white')
                draw = ImageDraw.Draw(img)
                text = page.extract_text() or f"Page {i+1}"
                draw.text((50, 50), text[:1000], fill='black')
                ip = out(f"page_{i+1}.jpg")
                img.save(str(ip), 'JPEG')
                zf.write(ip, f"page_{i+1}.jpg")
                ip.unlink(missing_ok=True)
        path.unlink(missing_ok=True)
        track_tool_use('PDF to JPG')
        track_tool_use('pdf2jpg')
        return send_file(zip_out, as_attachment=True, download_name="pdfnest-images.zip")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/jpg2pdf', methods=['POST'])
def jpg_to_pdf():
    try:
        from PIL import Image
        files = request.files.getlist('files')
        if not files: return jsonify({"error": "No files"}), 400
        saved = []
        images = []
        for f in files:
            p = save_upload(f); saved.append(p)
            img = Image.open(str(p)).convert('RGB')
            images.append(img)
        o = out("images.pdf")
        if images:
            images[0].save(str(o), save_all=True, append_images=images[1:])
        for p in saved: p.unlink(missing_ok=True)
        return send_file(o, as_attachment=True, download_name="pdfnest-images.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/rotate', methods=['POST'])
def rotate_pdf():
    try:
        file = request.files.get('file')
        angle = int(request.form.get('angle', 90))
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        writer = PdfWriter()
        for page in reader.pages:
            page.rotate(angle)
            writer.add_page(page)
        o = out("rotated.pdf")
        with open(o, 'wb') as f: writer.write(f)
        path.unlink(missing_ok=True)
        track_tool_use('Rotate PDF')
        track_tool_use('rotate')
        return send_file(o, as_attachment=True, download_name="pdfnest-rotated.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/protect', methods=['POST'])
def protect_pdf():
    try:
        file = request.files.get('file')
        password = request.form.get('password', '')
        if not file: return jsonify({"error": "No file"}), 400
        if not password: return jsonify({"error": "No password"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(password)
        o = out("protected.pdf")
        with open(o, 'wb') as f: writer.write(f)
        path.unlink(missing_ok=True)
        track_tool_use('Protect PDF')
        track_tool_use('protect')
        return send_file(o, as_attachment=True, download_name="pdfnest-protected.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/unlock', methods=['POST'])
def unlock_pdf():
    try:
        file = request.files.get('file')
        password = request.form.get('password', '')
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            reader.decrypt(password)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        o = out("unlocked.pdf")
        with open(o, 'wb') as f: writer.write(f)
        path.unlink(missing_ok=True)
        track_tool_use('Unlock PDF')
        track_tool_use('unlock')
        return send_file(o, as_attachment=True, download_name="pdfnest-unlocked.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/watermark', methods=['POST'])
def watermark_pdf():
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        file = request.files.get('file')
        text = request.form.get('text', 'CONFIDENTIAL')
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        wm_path = out("wm.pdf")
        c = canvas.Canvas(str(wm_path), pagesize=letter)
        c.setFont("Helvetica", 50)
        c.setFillAlpha(0.3)
        c.setFillColorRGB(0.5, 0.5, 0.5)
        c.saveState()
        c.translate(300, 400)
        c.rotate(45)
        c.drawCentredString(0, 0, text)
        c.restoreState()
        c.save()
        reader = PdfReader(str(path))
        wm_reader = PdfReader(str(wm_path))
        writer = PdfWriter()
        for page in reader.pages:
            page.merge_page(wm_reader.pages[0])
            writer.add_page(page)
        o = out("watermarked.pdf")
        with open(o, 'wb') as f: writer.write(f)
        path.unlink(missing_ok=True)
        wm_path.unlink(missing_ok=True)
        track_tool_use('Watermark PDF')
        track_tool_use('watermark')
        return send_file(o, as_attachment=True, download_name="pdfnest-watermarked.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/delete-pages', methods=['POST'])
def delete_pages():
    try:
        file = request.files.get('file')
        page_range = request.form.get('range', '')
        if not file: return jsonify({"error": "No file"}), 400
        path = save_upload(file)
        reader = PdfReader(str(path))
        total = len(reader.pages)
        pages_to_delete = set()
        if page_range:
            for part in page_range.split(','):
                part = part.strip()
                if '-' in part:
                    a,b = part.split('-')
                    pages_to_delete.update(range(int(a)-1, int(b)))
                else:
                    pages_to_delete.add(int(part)-1)
        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            if i not in pages_to_delete:
                writer.add_page(page)
        o = out("result.pdf")
        with open(o, 'wb') as f: writer.write(f)
        path.unlink(missing_ok=True)
        track_tool_use('Delete Pages')
        return send_file(o, as_attachment=True, download_name="pdfnest-result.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/qrcode', methods=['POST'])
def generate_qr():
    try:
        import qrcode
        data = request.form.get('text', '')
        if not data: return jsonify({"error": "No text"}), 400
        qr = qrcode.QRCode(version=1, box_size=10, border=4)
        qr.add_data(data); qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        o = out("qrcode.png"); img.save(str(o))
        return send_file(o, as_attachment=True, download_name="pdfnest-qrcode.png")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/admin/stats')
def admin_stats():
    stats = get_real_stats()
    return jsonify(stats)

import threading, time
def cleanup():
    while True:
        time.sleep(3600)
        for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
            for f in folder.iterdir():
                try:
                    if time.time() - f.stat().st_mtime > 7200: f.unlink()
                except: pass
threading.Thread(target=cleanup, daemon=True).start()

if __name__ == '__main__':
    print("\n========================================")
    print("  PDFnest is running!")
    print("  Open: http://localhost:5000")
    print("========================================\n")
    app.run(debug=False, port=5000)
